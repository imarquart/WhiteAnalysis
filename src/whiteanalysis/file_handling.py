from io import BytesIO

import tiktoken
from docx import Document
from pydantic import BaseModel
from pypdf import PdfReader
from unstructured.partition.pdf import partition_pdf


class PDFDocument(BaseModel):
    filename: str
    page: int
    text: str


def get_content_from_pdf(file: BytesIO, filename: str) -> list[PDFDocument]:
    """Loads a PDF file into a list of Document objects.

    Args:
        file: File object.
        filename: Name of the file.

    Returns:
        List of Document objects.
    """
    pdf_reader = PdfReader(file)
    text = ""
    pages: list[PDFDocument] = []
    for i, page in enumerate(pdf_reader.pages):
        text = page.extract_text()
        doc = PDFDocument(filename=filename, page=i, text=text)
        pages.append(doc)
    return pages


def get_content_from_unstructured(filename: str) -> list[PDFDocument]:
    """Loads a PDF file into a list of Document objects.

    Args:
        filename: Name of the file.

    Returns:
        List of Document objects.
    """
    elements = partition_pdf(filename)
    # Using tiktoken, cut into pages of 2048 tokens
    buffer = ""
    pages = []
    for element in elements:
        if len(buffer) + len(str(element)) > 2048:
            pages.append(buffer)
            buffer = ""
        buffer += str(element)
    if buffer:
        pages.append(buffer)
    documents = []
    for i, page in enumerate(pages):
        doc = PDFDocument(filename=filename, page=i, text=page)
        documents.append(doc)
    return documents


def get_content_from_docx(
    file: BytesIO, filename: str, encodings: tiktoken.Encoding, tokens_per_page: int
) -> list[PDFDocument]:
    """Extracts text from a DOCX file and splits it into pages based on token count.

    Args:
        file: DOCX file as bytes
        filename: Name of the file
        encodings: Tokenizer object with encode method
        tokens_per_page: Maximum tokens per page

    Returns:
        List of PDFDocument objects, each containing text within token limit
    """
    doc = Document(file)
    pages: list[PDFDocument] = []
    current_text: list[str] = []
    current_tokens = 0
    page_num = 0

    def create_page() -> None:
        nonlocal current_text, page_num, current_tokens
        if current_text:
            text = "\n".join(current_text)
            pages.append(PDFDocument(filename=filename, page=page_num, text=text))
            page_num += 1
            current_text = []
            current_tokens = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Check for section break
        if hasattr(paragraph._element, "sectPr"):
            create_page()
            continue

        # Check token count
        tokens = len(encodings.encode(text))
        if current_tokens + tokens > tokens_per_page:
            create_page()

        current_text.append(text)
        current_tokens += tokens

    # Add remaining text
    create_page()

    return pages or [PDFDocument(filename=filename, page=0, text="")]
