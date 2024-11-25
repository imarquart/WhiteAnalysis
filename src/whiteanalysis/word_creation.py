from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from whiteanalysis.prompts import Insights


def generate_word_report(
    responses: List[Insights],
    filename: str,
    case: str,
    model: str,
    output_path: str = "insights_report.docx",
):
    """
    Generate a minimally formatted Word document report from a list of Insights objects.
    Focuses on clear presentation of quotes for easy copying.

    Args:
        responses: List of Insights objects
        filename: Source filename
        case: Case description
        model: Model name
        output_path: Path where the Word file will be saved
    """
    doc = Document()
    # Set up default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header information in plain text
    doc.add_paragraph("CASES AND QUOTES")
    doc.add_paragraph(f"Source File: {filename}")
    doc.add_paragraph(f"Model: {model}")
    doc.add_paragraph("---")  # Simple separator

    # Process each insight
    for i, insight in enumerate(responses, 1):
        # Insight header
        doc.add_paragraph(f"\nINSIGHT SET {i}")

        # Context section
        doc.add_paragraph("GENERAL CONTEXT:")
        doc.add_paragraph(insight.general_context)
        doc.add_paragraph("RELEVANCE:")
        doc.add_paragraph(insight.general_relation)

        # Quotes section
        doc.add_paragraph("\nEXTRACTED QUOTES:")

        # Process each quote
        for j, quote in enumerate(insight.quotes, 1):
            # Add quote with clear separation
            doc.add_paragraph(f"\nQUOTE {j}:")
            doc.add_paragraph(f"Text: {quote.text}")
            doc.add_paragraph(f"Context: {quote.context}")
            doc.add_paragraph(f"Position: {quote.position}")
            doc.add_paragraph(
                f"Argument in draft: {quote.issue_in_draft if hasattr(quote, 'issue_in_draft') else ''}"
            )
            doc.add_paragraph(f"Relevance: {quote.relation}")

        doc.add_paragraph("\n" + "-" * 40)  # Section separator

    # Save the document
    doc.save(output_path)
    return output_path


def generate_word_report_pretty(
    responses: List[Insights],
    filename: str,
    case: str,
    model: str,
    output_path: str = "insights_report.docx",
):
    """
    Generate a Word document report from a list of Insights objects.

    Args:
        responses: List of Insights objects
        filename: Source filename
        case: Case description
        model: Model name
        output_path: Path where the Word file will be saved
    """
    doc = Document()

    # Set up default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header section
    title = doc.add_heading("Cases and Quotes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Meta information
    meta_info = doc.add_paragraph()
    meta_info.add_run("Source File: ").bold = True
    meta_info.add_run(filename)
    meta_info.add_run("\nModel: ").bold = True
    meta_info.add_run(model)
    meta_info.add_run("\nCase Description: ").bold = True
    meta_info.add_run(case)

    # Add spacing after header
    doc.add_paragraph()

    # Process each insight
    for i, insight in enumerate(responses, 1):
        # Insight header
        doc.add_heading(f"Insight Set {i}", level=1)

        # Context section with light gray background
        context_table = doc.add_table(rows=1, cols=1)
        context_table.style = "Table Grid"
        cell = context_table.rows[0].cells[0]

        # General Context
        context_para = cell.paragraphs[0]
        context_para.add_run("General Context\n").bold = True
        context_para.add_run(insight.general_context)
        context_para.add_run("\n\nRelevance\n").bold = True
        context_para.add_run(insight.general_relation)

        # Add spacing
        doc.add_paragraph()

        # Quotes section
        doc.add_heading("Extracted Quotes", level=2)

        # Process each quote
        for quote in insight.quotes:
            # Create a table for each quote to maintain consistent formatting
            quote_table = doc.add_table(rows=1, cols=1)
            quote_table.style = "Table Grid"
            quote_cell = quote_table.rows[0].cells[0]

            # Quote text
            quote_para = quote_cell.paragraphs[0]
            quote_para.add_run(quote.text).bold = True

            # Quote context
            quote_cell.add_paragraph(quote.context)

            # Position
            position_para = quote_cell.add_paragraph()
            position_para.add_run("Position: ").bold = True
            position_para.add_run(quote.position)

            # Relevance with light blue background
            relevance_table = quote_cell.add_table(rows=1, cols=1)
            relevance_cell = relevance_table.rows[0].cells[0]
            relevance_para = relevance_cell.paragraphs[0]
            relevance_para.add_run("Relevance: ").bold = True
            relevance_para.add_run(quote.relation)

            # Add spacing between quotes
            doc.add_paragraph()

        # Add spacing between insight sets
        doc.add_paragraph()

    # Save the document
    doc.save(output_path)
    return output_path
